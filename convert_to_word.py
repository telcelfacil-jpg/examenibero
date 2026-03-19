import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('> [!'):
            # Simple handling for GitHub alerts
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            # Basic text with bold/italic handling if needed, but keeping it simple for now
            # as python-docx requires more logic for inline formatting from markdown
            doc.add_paragraph(line.replace('**', '').replace('*', ''))

    doc.save(docx_path)

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python convert.py <input_md> <output_docx>")
    else:
        markdown_to_docx(sys.argv[1], sys.argv[2])
